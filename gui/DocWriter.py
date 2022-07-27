from docx import *
from docx.shared import Cm
from docx.enum.style import WD_STYLE_TYPE
from docx import __version__

print(__version__)


data = {
    "ime" : "Alessandra Melesi",
    "jmbg" : 142535262134,
    "lbo" : 252341423,
    "filijala" : "Tukisha Markovic",
    "knjizica" : "25.02.2023",
    "vid_terapije" : 3,
    "lek_lista" : [2, 5],
    "rezim_doziranja" : 0,
    "godine_zivota" : 27,
    "pol" : "M",
    "tm" : 75,
    "tv" : 182,
    "bmi" : 21,
    "lokalizacija" : [2, 4, 5],
    "ponasanje_bolesti" : [3],
    "fistula" : [],
    "apsces" : [2],
    "trajanje_bolesti" : "6 meseci",
    "ekstraintestinalne_manifestacije" : True,
    "vrsta_eim" : "Brate ne znam",
    "klinicka_aktivnost" : 4,
    "crp" : 10.2,
    "kalprotektin" : "???",
    "quantiferon_test" : True,
    "ppd3" : True,
    "rtg_uredan" : True,
    "clostridium" : False,
    "hbs_antigen" : True,
    "endoskopija" : "24.05.2023.",
    "prisustvo_ulceracija" : True,
    "operacija" : True,
    "br_operacija" : 2,
    "razlog_operacije" : [],
    "prisutna" : [],

    "kortiko" : True,
    "kortiko_oralno" : 2,
    "kortiko_doza" : "2mg",
    "kortiko_pocetak" : '24.7.2022.',
    "kortiko_kraj" : True,
    "kortiko_poboljsanje" : True,
    "kortiko_kortikozavisan" : True,
    "kortiko_kortikorezistentan" : True,

    "aza" : True,
    "mtx" : True,
    "aza_doza" : "2mg",
    "aza_pocetak" : '24.7.2022.',
    "aza_kraj" : True,
    "aza_alergija" : True,
    "aza_rezistentan" : True,

    "ifx" : True,
    "ada" : True,
    "vdz" : True,
    "ifx_doza" : "2mg",
    "ifx_pocetak" : '24.7.2022.',
    "ifx_broj_ampula" : 7,
    "ifx_duzina_terapije" : "6 meseci",
    "ifx_kraj" : True,
    "ifx_alergija" : False,
    
    "nuspojave_na_lekove" : False,
    "nuspojave_lek" : "",
    "nuspojave_vrsta" : "",
    "trajanje_biloske_terapije" : "9 meseci",

    "datum" : "26.7.2022."

}


#def docWriting(data):
doc = Document()
doc.styles['Normal'].font.name = 'Arial'
doc.styles['Normal'].paragraph_format.line_spacing = 1.0
doc.styles['Normal'].paragraph_format.space_after = Cm(0)

myStyle = doc.styles.add_style('L1', WD_STYLE_TYPE.PARAGRAPH)
myStyle.base_style = doc.styles['List Number']
myStyle.paragraph_format.left_indent = Cm(1.27)
myStyle.paragraph_format.space_before = Cm(0)
myStyle.paragraph_format.space_after = Cm(0)
myStyle.paragraph_format.line_spacing = 1.15

myStyle = doc.styles.add_style('L1_Selected', WD_STYLE_TYPE.PARAGRAPH)
myStyle.base_style = doc.styles['L1']
myStyle.paragraph_format.left_indent = Cm(1.27)
myStyle.paragraph_format.space_before = Cm(0)
myStyle.paragraph_format.space_after = Cm(0)
myStyle.paragraph_format.line_spacing = 1.15
myStyle.font.bold = True
myStyle.font.underline = True

myStyle = doc.styles.add_style('L2', WD_STYLE_TYPE.PARAGRAPH)
myStyle.base_style = doc.styles['List Number']
myStyle.paragraph_format.left_indent = Cm(1.27)
myStyle.paragraph_format.space_before = Cm(0)
myStyle.paragraph_format.space_after = Cm(0)
myStyle.paragraph_format.line_spacing = 1.15

myStyle = doc.styles.add_style('L2_Selected', WD_STYLE_TYPE.PARAGRAPH)
myStyle.base_style = doc.styles['L2']
myStyle.paragraph_format.left_indent = Cm(1.27)
myStyle.paragraph_format.space_before = Cm(0)
myStyle.paragraph_format.space_after = Cm(0)
myStyle.paragraph_format.line_spacing = 1.15
myStyle.font.bold = True
myStyle.font.underline = True


#margins
doc.sections[0].top_margin = Cm(1.27)
doc.sections[0].bottom_margin = Cm(1.27)
doc.sections[0].left_margin = Cm(2.54)
doc.sections[0].right_margin = Cm(2.54)

#title
doc.add_paragraph().add_run('\nFORMULAR SA KONZILIJARNIM MIŠLJENJEM ZA PRIMENU BIOLOŠKE TERAPIJE KOD PACIJENATA OBOLELIH OD CROHN-OVE BOLESTI\n\n').bold = True

# 1. Cluster
#IME, JMBG, LBO, FILIJALA, KNJIZICA, USTANOVA
para = doc.add_paragraph()
para.add_run(f'Ime i prezime osiguranog lica: {data["ime"]}\n').bold = True
lin = para.add_run(f'JMBG: {data["jmbg"]}\t\t')
lin.add_text(f'LBO: {data["lbo"]}\t')
lin.add_text(f'Filijala: {data["filijala"]}')
para.add_run(f'\nDatum važenja zdravstvene knjižice: {data["knjizica"]}\n')
para.add_run('Ustanova/Referentni Centar: ') #QUESTION TODO da li mora da se ukucava u polje?

#para.add_run('Klinika za gastroenterohepatologiju UKCS').bold = True

# 2. Cluster
#VID TERAPIJE, LEK, REZIM DOZIRANJA
para = doc.add_paragraph()
para.add_run('\n1. Uvođenje leka u terapiju  ').bold = (data["vid_terapije"] == 1)
para.add_run('2. Terapija održavanja  ').bold = (data["vid_terapije"] == 2)
para.add_run('3. Promena leka   ').bold = (data["vid_terapije"] == 3)
para.add_run('4. Prekid terapije ').bold = (data["vid_terapije"] == 4)

#LEK
para = doc.add_paragraph()
para.add_run('Lek: ').bold = True
doc.add_paragraph('Remicade', style = 'L1_Selected' if (1 in data["lek_lista"]) else 'L1')
doc.add_paragraph('Remsima', style = 'L1_Selected' if (2 in data["lek_lista"]) else 'L1')
doc.add_paragraph('Inflectra', style = 'L1_Selected' if (3 in data["lek_lista"]) else 'L1')
doc.add_paragraph('Humira', style = 'L1_Selected' if (4 in data["lek_lista"]) else 'L1')
doc.add_paragraph('Amgevita', style = 'L1_Selected' if (5 in data["lek_lista"]) else 'L1')
doc.add_paragraph('Hyrimoz', style = 'L1_Selected' if (6 in data["lek_lista"]) else 'L1')
doc.add_paragraph('Idacio', style = 'L1_Selected' if (7 in data["lek_lista"]) else 'L1')
doc.add_paragraph('Entyvio', style = 'L1_Selected' if (8 in data["lek_lista"]) else 'L1')


para = doc.add_paragraph()
para.add_run('\nRežim doziranja: ')
para.add_run('1. STANDARDNI\t\t').bold = (data["rezim_doziranja"] == 1)
para.add_run('2. OPTIMIZACIJA').bold = (data["rezim_doziranja"] == 2)


# 3. Cluster
para = doc.add_paragraph()
para.add_run('\n\nPol: ')
para.add_run('Ž ').bold = (data["pol"] != "M")
para.add_run('/')
para.add_run(' M').bold = (data["pol"] == "M")

para.add_run(f'\nGodine života: {data["godine_zivota"]}')

para.add_run(f'\nTM: {data["tm"]}')
#para.add_run(f'TV= {data["tv"]}\t')
para.add_run(f'\nBMI= {data["bmi"]}')


para.add_run('\n\nDg: ').bold = True
para.add_run('Mb Crohn (K 50)')


# Lokalizacija list
para = doc.add_paragraph()
para.add_run('\nLokalizacija: ').bold = True
para.add_run('\n\t1. Ileum').bold = (1 in data["lokalizacija"])
para.add_run('\n\t2. Ileum + Kolon').bold = (2 in data["lokalizacija"])
para.add_run('\n\t3. Kolon').bold = (3 in data["lokalizacija"])
para.add_run('\n\t4. Rektum').bold = (4 in data["lokalizacija"])
para.add_run('\n\t5. Proksimalni segmenti GIT-a').bold = (5 in data["lokalizacija"])
#para.add_run('\n\t6. Drugo').bold = (6 in data["lokalizacija"])

#TODO - dodaj opciju za DRUGO 



#Ponasanje bolesti cluster
para = doc.add_paragraph()
para.add_run('\nPonašanje bolesti: ').bold = True
para.add_run('1. Inflamatorna forma ').bold = (1 in data["ponasanje_bolesti"])
para.add_run('2. Penetrantna  ').bold = (2 in data["ponasanje_bolesti"])
para.add_run('3. Stenozantna').bold = (3 in data["ponasanje_bolesti"])


#Fistula cluster
para = doc.add_paragraph()
para.add_run('\nFistula/e: ').bold = True
para.add_run('Da').bold = (len(data["fistula"]) > 0)
para.add_run('/')
para.add_run('Ne ').bold = (len(data["fistula"]) == 0)
para.add_run('\n\t1. Perianalna').bold = (1 in data["fistula"])
para.add_run('\n\t2. Enterokutana').bold = (2 in data["fistula"])
para.add_run('\n\t3. Enterovaginalna').bold = (3 in data["fistula"])
para.add_run('\n\t4. Enteroenteralna').bold = (4 in data["fistula"])
#TODO - dodaj opciju za DRUGO 

#Apsces cluster
para = doc.add_paragraph()
para.add_run('\nApces: ').bold = True
para.add_run('Da').bold = (len(data["apsces"]) > 0)
para.add_run('/')
para.add_run('Ne ').bold = (len(data["apsces"]) == 0)
para.add_run('\n\t1. Perianalni').bold = (1 in data["apsces"])
para.add_run('\n\t2. Interintestinalni').bold = (2 in data["apsces"])
#para.add_run('\n3. Drugo').bold = (3 in data["apsces"])
#TODO - dodaj opciju za DRUGO


# 4. Cluster
#DUZINA TRAJANJA, MANIFESTACIJE, VRSTA EIM
para = doc.add_paragraph()
para.add_run(f'\nDužina trajanja bolesti: {data["trajanje_bolesti"]}').bold = True
para.add_run('\nEkstraintestinalne manifestacije: ')
para.add_run('Da').bold = data["ekstraintestinalne_manifestacije"] == False
para.add_run('/')
para.add_run('Ne ').bold = data["ekstraintestinalne_manifestacije"] == True
para.add_run(f'\nVrsta EIM: {data["vrsta_eim"]}')
doc.add_page_break()

#Klinicka aktivnost bolesti cluster
para = doc.add_paragraph()
para.add_run('Klinička aktivnost bolesti: ').bold = True
para.add_run('1. Blaga ').bold = (data["klinicka_aktivnost"] == 1)
para.add_run('2. Umerena ').bold = (data["klinicka_aktivnost"] == 2)
para.add_run('3. Izrazito aktivna ').bold = (data["klinicka_aktivnost"] == 3)
para.add_run('4. Remisija').bold = (data["klinicka_aktivnost"] == 4)


# 5. Cluster
#CRP, FEKALNI KALPROTEKTIN, QUANTIFERON, RTG, CLOSTRIDIUM, HBs
para = doc.add_paragraph()
para.add_run(f'\nVrednost CRP-a: {data["crp"]}')
para.add_run(f'\nFekalni kalprotektin: {data["kalprotektin"]}')

para.add_run('\nQuantiferon test: ')
para.add_run('1. poz  ').bold = data["quantiferon_test"] == True
para.add_run('2. neg').bold = data["quantiferon_test"] == False

para.add_run('\nPPD 3 test: ')
para.add_run('1. poz  ').bold = data["ppd3"] == True
para.add_run('2. neg').bold = data["ppd3"] == False

para.add_run('\nRtg srca i pluća: ')
para.add_run('1. Nalaz uredan  ').bold = data["rtg_uredan"] == True
para.add_run('2. Izmenjen').bold = data["rtg_uredan"] == False

para.add_run('\nTest na Clostridium difficile: ')
para.add_run('1. poz  ').bold = data["clostridium"] == True
para.add_run('2. neg').bold = data["clostridium"] == False

para.add_run('\nHBs antigen: ')
para.add_run('1. poz  ').bold = data["hbs_antigen"] == True
para.add_run('2. neg').bold = data["hbs_antigen"] == False



#6. Cluster
#ENDOSKOPIJA, OPERACIJA, BR OPERACIJA, RAZLOG OPERACIJE, PRISUTNA
para = doc.add_paragraph()

para.add_run('\nEndoskopija: ').bold = True
para.add_run('1. Prisustvo ulceracija  ').bold = data["prisustvo_ulceracija"] == True
para.add_run('2. Odsustvo ulceracija').bold = data["prisustvo_ulceracija"] == False
#para.add_run(f'({data["endoskopija"]})').bold = data["endoskopija"] > 0

para.add_run('\n\nOperacije: ').bold = True
para.add_run('DA').bold = data["operacija"] == True
para.add_run('/')
para.add_run('NE').bold = data["operacija"] == False

para.add_run('\nBroj operacija: ')
para.add_run('1. Jedna ').bold = data["br_operacija"] == 1
para.add_run('2. >= 2').bold = data["br_operacija"] >= 2

para.add_run('\nRazlog operacije: ')
para.add_run('1. Stenoza \t').bold = (1 in data["razlog_operacije"] )
para.add_run('2. Penetrantna bolest \t').bold = (2 in data["razlog_operacije"])
para.add_run('3. Drugi').bold = (3 in data["razlog_operacije"])
#TODO dodaj Drugi

para.add_run('\nPrisutna: ')
para.add_run('Ileostoma').bold = (1 in data["prisutna"])
para.add_run('/')
para.add_run('Kolostoma').bold = (2 in data["prisutna"])


#DOSADASNJA TERAPIJA
para = doc.add_paragraph()
para.add_run('\nDosadašnja terapija: ')

# 1. KKS
para = doc.add_paragraph()
para.add_run('\n1. ')
para.add_run('KKS per os/i.v.: ').bold = data["kortiko"] == True
para.add_run(f'\n    Doza: {data["kortiko_doza"]}')
para.add_run(f'\n    Početak terapije: {data["kortiko_pocetak"]}')

para.add_run('\n    Kraj terapije').bold = data["kortiko_kraj"] == True
para.add_run('/')
para.add_run('Još uvek traje').bold = data["kortiko_kraj"] == False

para.add_run('\n    Remisija/Poboljšanje trećeg dana i.v. KKS: ')
para.add_run('1. Da ').bold = data["kortiko_poboljsanje"] == True
para.add_run('2. Ne').bold = data["kortiko_poboljsanje"] == False

para.add_run('\n    Kortiko zavisan: ')
para.add_run('1. Da ').bold = data["kortiko_kortikozavisan"] == True
para.add_run('\t2. Ne').bold = data["kortiko_kortikozavisan"] == False

para.add_run('\n    Kortiko rezistentan: ')
para.add_run('1. Da ').bold = data["kortiko_kortikorezistentan"] == True
para.add_run('\t2. Ne').bold = data["kortiko_kortikorezistentan"] == False


# 2. AZA
para = doc.add_paragraph()
para.add_run('\n2. ')
para.add_run('AZA').bold = data["aza"] == True
para.add_run('/')
para.add_run('MTX').bold = data["mtx"] == True
para.add_run(':')
para.add_run(f'\n    Doza: {data["aza_doza"]}')
para.add_run(f'\n    Početak terapije: {data["aza_pocetak"]}')

para.add_run('\n    Kraj terapije').bold = data["aza_kraj"] == True
para.add_run('/')
para.add_run('Još uvek traje').bold = data["aza_kraj"] == False

para.add_run('\n    Alergija na lek: ')
para.add_run('1. Da ').bold = data["aza_alergija"] == True
para.add_run('\t2. Ne').bold = data["aza_alergija"] == False

para.add_run('\n    Rezistentan na imunosupresive: ')
para.add_run('1. Da ').bold = data["aza_rezistentan"] == True
para.add_run('\t2. Ne').bold = data["aza_rezistentan"] == False


#3. IFX/ADA/VDZ
para = doc.add_paragraph()
para.add_run('\n3. ')
para.add_run('IFX').bold = data["ifx"] == True
para.add_run('/')
para.add_run('ADA').bold = data["ada"] == True
para.add_run('/')
para.add_run('VDZ').bold = data["vdz"] == True
para.add_run(':')
para.add_run(f'\n    Doza: {data["ifx_doza"]}')
para.add_run(f'\n    Početak terapije: {data["ifx_pocetak"]}')
para.add_run(f'\n    Broj ampula po ciklusu: {data["ifx_broj_ampula"]}')
para.add_run(f'\n    Dužina terapije do sada: {data["ifx_duzina_terapije"]}')

para.add_run('\n    Kraj terapije').bold = data["ifx_kraj"] == True
para.add_run('/')
para.add_run('Još uvek traje').bold = data["ifx_kraj"] == False

para.add_run('\n    Alergična na lek: ')
para.add_run('1. Da ').bold = data["ifx_alergija"] == True
para.add_run('\t2. Ne').bold = data["ifx_alergija"] == False



# Pred kraj
para = doc.add_paragraph()
para.add_run('\nNus pojave na dosadašnje lekove: ').bold = True
para.add_run('1.Ne ').bold = data["nuspojave_na_lekove"] == False
para.add_run('2.Da').bold = data["nuspojave_na_lekove"] == True

para.add_run(f'\nLek: {data["nuspojave_lek"]}')
para.add_run(f'\nVrsta manifestacija: {data["nuspojave_vrsta"]}')


# KOMENTAR
para = doc.add_paragraph()
para.add_run('\n\nUKUPNA DUŽINA BIOLOŠKE TERAPIJE DO SADA:').bold = True

doc.add_page_break()


# Kraj
para = doc.add_paragraph()

para.add_run('KONZILIJARNO MIŠLJENJE:').bold = True

for i in range (25):
    para.add_run('\n')
    
para.add_run(f'Datum: {data["datum"]}')
para.add_run('\t\t\t\tPotpis članova komisije KC/KBC:')

#save the document
doc.save('output.docx')